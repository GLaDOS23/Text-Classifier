#!pip install requests
#!pip install pdfplumber python-docx

import os
import pdfplumber
import docx
from docx import Document
import requests


from google.colab import drive
drive.mount('/content/drive', force_remount=True)

# Яндекс GPT
def YaGPT(input_str):
  text = "Составь из этого текста вопрос для теста и один ответ на него (без дополнительной информации): " + str(input_str)
  prompt = {
      "modelUri": "gpt://b1g1o9j4lup4tvhfe5u6/yandexgpt-lite",
      "completionOptions": {
          "stream": False,
          "temperature": 0.6,
          "maxTokens": "2000"
      },
      "messages": [
          {
              "role": "user",
              "text": text
          }
      ]
      
  }


  url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
  headers = {
      "Content-Type": "application/json",
      "Authorization": "Api-Key AQVN36QRntV3Gys6tAvWaL1nEZeEYkCkW2hsYdqV"
  }


  response = requests.post(url, headers=headers, json=prompt)

  result = response.text
  #print(result)
  result = result.replace('{"result":{"alternatives":[{"message":{"role":"assistant","text":"','')
  result = result[:-156]
  return result


# Найти жирный и подчёркнутый текст
def find_bold_and_underline_text(docx_file):
  doc = docx.Document(docx_file)
  find_text = []

  for paragraph in doc.paragraphs:
      for run in paragraph.runs:
          if run.bold or run.underline:
              find_text.append(run.text)
          #if run.underline:
              #underline_text.append(run.text)

  return find_text



# Поиск ключевого слова и копирование предложения
def find_and_extract_text(docx_file, search_word):
    doc = docx.Document(docx_file)
    extracted_texts = []

    for paragraph in doc.paragraphs:
        if search_word in paragraph.text:
            # Найти индекс слова в абзаце
            index = paragraph.text.find(search_word)

            # Извлечь текст рядом со словом
            extracted_text = ""
            while index < len(paragraph.text) and paragraph.text[index] not in [".", "\n"]:
                index += 1
                extracted_text += paragraph.text[index]

            extracted_texts.append(extracted_text)

    return extracted_texts


# Проверка расширения файла и конвертация .pdf в .docx
def check_file_format(file_path):
    new_file = None
    # Получить расширение файла
    file_extension = os.path.splitext(file_path)[1]

    # Определить формат файла на основе расширения
    if file_extension == ".docx":
        new_file = file_path

    elif file_extension == ".pdf":
      # Получить папку расположения файла
      file_dir = os.path.dirname(file_path)
      print("Переводим файл pdf в docx и сохраняем")
      # Переводим .pdf в .docx
      pdf = pdfplumber.open(file_path)
      doc = Document()
      for page in pdf.pages:
        text = page.extract_text()
        doc.add_paragraph(text)
      # Сохраняем
      new_file = os.path.join(file_dir, f"new_file.docx")
      doc.save(new_file)
      #doc.save('/content/drive/My Drive/my_doc.docx')


    else:
        print("Неверный формат файла")

    return new_file

while True:

  file_path = input("Введите путь к файлу: ")

  # Проверить, существует ли файл
  if os.path.exists(file_path):
    # Проверка расширения файла
    new_file = check_file_format (file_path)
    if new_file != None:
      # Поиск текста
      find_themes = find_bold_and_underline_text(new_file)
      print("Найденные темы:")
      for i in find_themes:
        print(i)
      # Составляем вопросы по темам
      while True:
        inputs = input("Введите тему:")
        if inputs == "стоп":
          break
        else:
          find_text = find_and_extract_text(new_file, inputs)
          # Проверка на наличие нужного текста
          if len(find_text) != 0:
            resulting_text = YaGPT(find_text)
            print(resulting_text)
          else:
            print("Тема не найдена")
    

  else:
    print("Файл не существует.")
#file = "/content/drive/My Drive/lecture_1_3.docx"
